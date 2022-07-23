from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtWidgets import QFileDialog
from PIL import Image
import pytesseract
import docx
from docxcompose.composer import Composer
from docx.shared import Pt, Mm
import os


class img2doc:
	
	def __init__(self, dir_path, to_path, font='Times New Roman', font_size=22, page_height=350, language='english', save_formats=['docx']):
		
		if type(save_formats) == str : save_formats = [save_formats]
		
		languages = {"english":"eng",
					 "arabic":"ara",
					 "french":"fra"}

		self.language = languages[language.lower()]
		
		self.dir_path = dir_path
		try : self.images_names = sorted([int(c.split('.')[0]) for c in os.listdir(self.dir_path)])
		except : self.images_names = sorted([c.split('.')[0] for c in os.listdir(self.dir_path)])
		self.images_names = [str(c) for c in self.images_names]
		self.images_names_full = [str(c)+'.jpg' for c in self.images_names]
		self.images_paths = [os.path.join(self.dir_path, img) for img in self.images_names_full]
		self.to_path = to_path
		self.font = font
		self.font_size = font_size
		self.page_height = page_height
		self.documents = []
		self.save_formats = save_formats
		self.time_type = 'seconds'
		self.predicted_time = (2 * len(self.images_paths))
		if self.predicted_time >= 60:
			self.predicted_time = self.predicted_time // 60 
			self.time_type = 'minute' if self.predicted_time == 1 else 'minutes'  
		for _ in range(len(self.save_formats)):
			if self.save_formats[_][0] != '.' : self.save_formats[_] = '.' + self.save_formats[_] 
		
		
	def convert(self):
		for image_path in self.images_paths:
			try:	
				text = str(pytesseract.image_to_string(Image.open(image_path), lang=self.language).rstrip())
				doc = docx.Document()
				doc.styles['Normal'].font.name = self.font
				doc.styles['Normal'].font.size = Pt(self.font_size)
				doc.sections[0].page_height = Mm(self.page_height)
				doc.add_paragraph(text)
				self.documents.append(doc)
			except:
				pass
			
	def merge(self):
		master = docx.Document()
		master.styles['Normal'].font.name = self.font
		master.styles['Normal'].font.size = Pt(self.font_size)
		master.sections[0].page_height = Mm(self.page_height)
		composer = Composer(master)
		
		for doc_temp in self.documents:
			composer.append(doc_temp)
			
		for save_format in self.save_formats:	
			composer.save(os.path.join(self.to_path, 'img2doc'+save_format))

class Ui_MainWindow(object):
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(735, 600)
        MainWindow.setMinimumSize(QtCore.QSize(735, 576))
        MainWindow.setMaximumSize(QtCore.QSize(735, 600))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setBold(False)
        font.setItalic(False)
        font.setUnderline(False)
        font.setWeight(50)
        font.setStrikeOut(False)
        font.setKerning(True)
        MainWindow.setFont(font)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.label = QtWidgets.QLabel(self.centralwidget)
        self.label.setGeometry(QtCore.QRect(50, 10, 641, 91))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(60)
        font.setBold(False)
        font.setItalic(False)
        font.setUnderline(False)
        font.setWeight(50)
        font.setStrikeOut(False)
        font.setKerning(True)
        self.label.setFont(font)
        self.label.setObjectName("label")
        self.label_2 = QtWidgets.QLabel(self.centralwidget)
        self.label_2.setGeometry(QtCore.QRect(100, 110, 521, 17))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(11)
        font.setBold(False)
        font.setItalic(False)
        font.setUnderline(False)
        font.setWeight(50)
        font.setStrikeOut(False)
        font.setKerning(True)
        self.label_2.setFont(font)
        self.label_2.setObjectName("label_2")
        self.label_3 = QtWidgets.QLabel(self.centralwidget)
        self.label_3.setGeometry(QtCore.QRect(250, 550, 221, 16))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(11)
        font.setBold(False)
        font.setItalic(False)
        font.setUnderline(False)
        font.setWeight(50)
        font.setStrikeOut(False)
        font.setKerning(True)
        self.label_3.setFont(font)
        self.label_3.setObjectName("label_3")
        self.label_4 = QtWidgets.QLabel(self.centralwidget)
        self.label_4.setGeometry(QtCore.QRect(60, 180, 101, 31))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(20)
        font.setBold(False)
        font.setItalic(False)
        font.setUnderline(False)
        font.setWeight(50)
        font.setStrikeOut(False)
        font.setKerning(True)
        self.label_4.setFont(font)
        self.label_4.setObjectName("label_4")
        self.pathOption = QtWidgets.QLineEdit(self.centralwidget)
        self.pathOption.setGeometry(QtCore.QRect(200, 180, 301, 31))
        self.pathOption.setObjectName("pathOption")
        self.pathBrowse = QtWidgets.QPushButton(self.centralwidget)
        self.pathBrowse.setGeometry(QtCore.QRect(520, 180, 161, 31))
        self.pathBrowse.setObjectName("pathBrowse")
        self.pathBrowse.clicked.connect(self.oca_path)
        self.label_5 = QtWidgets.QLabel(self.centralwidget)
        self.label_5.setGeometry(QtCore.QRect(60, 250, 91, 31))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(20)
        font.setBold(False)
        font.setItalic(False)
        font.setUnderline(False)
        font.setWeight(50)
        font.setStrikeOut(False)
        font.setKerning(True)
        self.label_5.setFont(font)
        self.label_5.setScaledContents(False)
        self.label_5.setObjectName("label_5")
        self.saveToOption = QtWidgets.QLineEdit(self.centralwidget)
        self.saveToOption.setGeometry(QtCore.QRect(200, 250, 301, 31))
        self.saveToOption.setObjectName("saveToOption")
        self.label_6 = QtWidgets.QLabel(self.centralwidget)
        self.label_6.setGeometry(QtCore.QRect(60, 400, 101, 31))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(20)
        font.setBold(False)
        font.setItalic(False)
        font.setUnderline(False)
        font.setWeight(50)
        font.setStrikeOut(False)
        font.setKerning(True)
        self.label_6.setFont(font)
        self.label_6.setScaledContents(False)
        self.label_6.setObjectName("label_6")
        self.pdfOption = QtWidgets.QCheckBox(self.centralwidget)
        self.pdfOption.setGeometry(QtCore.QRect(200, 400, 61, 31))
        font = QtGui.QFont()
        font.setPointSize(14)
        self.pdfOption.setFont(font)
        self.pdfOption.setObjectName("pdfOption")
        self.docxOption = QtWidgets.QCheckBox(self.centralwidget)
        self.docxOption.setGeometry(QtCore.QRect(270, 400, 81, 31))
        font = QtGui.QFont()
        font.setPointSize(14)
        self.docxOption.setFont(font)
        self.docxOption.setObjectName("docxOption")
        self.docOption = QtWidgets.QCheckBox(self.centralwidget)
        self.docOption.setGeometry(QtCore.QRect(350, 400, 81, 31))
        font = QtGui.QFont()
        font.setPointSize(14)
        self.docOption.setFont(font)
        self.docOption.setObjectName("docOption")
        self.saveToBrowse = QtWidgets.QPushButton(self.centralwidget)
        self.saveToBrowse.setGeometry(QtCore.QRect(520, 250, 161, 31))
        self.saveToBrowse.setObjectName("saveToBrowse")
        self.saveToBrowse.clicked.connect(self.oca_save)
        self.label_7 = QtWidgets.QLabel(self.centralwidget)
        self.label_7.setGeometry(QtCore.QRect(60, 330, 131, 31))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(20)
        font.setBold(False)
        font.setItalic(False)
        font.setUnderline(False)
        font.setWeight(50)
        font.setStrikeOut(False)
        font.setKerning(True)
        self.label_7.setFont(font)
        self.label_7.setObjectName("label_7")
        self.languageOption = QtWidgets.QComboBox(self.centralwidget)
        self.languageOption.setGeometry(QtCore.QRect(200, 330, 101, 31))
        self.languageOption.setObjectName("languageOption")
        self.languageOption.addItem("")
        self.languageOption.addItem("")
        self.languageOption.addItem("")
        self.label_8 = QtWidgets.QLabel(self.centralwidget)
        self.label_8.setGeometry(QtCore.QRect(370, 330, 71, 31))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(20)
        font.setBold(False)
        font.setItalic(False)
        font.setUnderline(False)
        font.setWeight(50)
        font.setStrikeOut(False)
        font.setKerning(True)
        self.label_8.setFont(font)
        self.label_8.setObjectName("label_8")
        self.fontOption = QtWidgets.QFontComboBox(self.centralwidget)
        self.fontOption.setGeometry(QtCore.QRect(460, 330, 221, 31))
        self.fontOption.setObjectName("fontOption")
        self.pushButton_3 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_3.setGeometry(QtCore.QRect(58, 450, 621, 91))
        
        font = QtGui.QFont()
        font.setPointSize(40)
        self.pushButton_3.setFont(font)
        self.pushButton_3.setObjectName("pushButton_3")
        self.pushButton_3.clicked.connect(self.oca_convert)
        self.docOption_2 = QtWidgets.QCheckBox(self.centralwidget)
        self.docOption_2.setGeometry(QtCore.QRect(440, 400, 81, 31))
        font = QtGui.QFont()
        font.setPointSize(14)
        self.docOption_2.setFont(font)
        self.docOption_2.setObjectName("docOption_2")
        self.docOption_3 = QtWidgets.QCheckBox(self.centralwidget)
        self.docOption_3.setGeometry(QtCore.QRect(510, 400, 81, 31))
        font = QtGui.QFont()
        font.setPointSize(14)
        self.docOption_3.setFont(font)
        self.docOption_3.setObjectName("docOption_3")
        self.label_9 = QtWidgets.QLabel(self.centralwidget)
        self.label_9.setGeometry(QtCore.QRect(120, 140, 471, 20))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.label_9.setFont(font)
        self.label_9.setObjectName("label_9")
        self.label.raise_()
        self.label_2.raise_()
        self.label_4.raise_()
        self.pathOption.raise_()
        self.pathBrowse.raise_()
        self.label_5.raise_()
        self.saveToOption.raise_()
        self.label_6.raise_()
        self.pdfOption.raise_()
        self.docxOption.raise_()
        self.docOption.raise_()
        self.saveToBrowse.raise_()
        self.label_7.raise_()
        self.languageOption.raise_()
        self.label_8.raise_()
        self.fontOption.raise_()
        self.pushButton_3.raise_()
        self.docOption_2.raise_()
        self.docOption_3.raise_()
        self.label_3.raise_()
        self.label_9.raise_()
        MainWindow.setCentralWidget(self.centralwidget)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "Img2doc"))
        self.label.setText(_translate("MainWindow", "Image2Document"))
        self.label_2.setText(_translate("MainWindow", "Convert images containing texts into a PDF/DOC/DOCX/PPT/PPTX document"))
        self.label_3.setText(_translate("MainWindow", "-Developer: Dhia Eddine Hanafi-"))
        self.label_4.setText(_translate("MainWindow", "Images"))
        self.pathBrowse.setText(_translate("MainWindow", "Browse"))
        self.label_5.setText(_translate("MainWindow", "Save to"))
        self.label_6.setText(_translate("MainWindow", "Save as"))
        self.pdfOption.setText(_translate("MainWindow", "PDF"))
        self.docxOption.setText(_translate("MainWindow", "DOC"))
        self.docOption.setText(_translate("MainWindow", "DOCX"))
        self.saveToBrowse.setText(_translate("MainWindow", "Browse"))
        self.label_7.setText(_translate("MainWindow", "Language"))
        self.languageOption.setItemText(0, _translate("MainWindow", "English"))
        self.languageOption.setItemText(1, _translate("MainWindow", "Arabic"))
        self.languageOption.setItemText(2, _translate("MainWindow", "French"))
        self.label_8.setText(_translate("MainWindow", "Font"))
        self.pushButton_3.setText(_translate("MainWindow", "Convert"))
        self.docOption_2.setText(_translate("MainWindow", "PPT"))
        self.docOption_3.setText(_translate("MainWindow", "PPTX"))
        self.label_9.setText(_translate("MainWindow", "استخراج النصوص من الصور وحفظها في ملف PDF/DOC/DOCX/PPT/PPTX"))
        
    def oca_convert(self):
        font = self.fontOption.currentText()
        language = self.languageOption.currentText()
        path = self.pathOption.text()
        save_to = self.saveToOption.text()
        save_formats = []
        if self.pdfOption.isChecked()   : save_formats.append('pdf')
        if self.docOption.isChecked()   : save_formats.append('docx')
        if self.docxOption.isChecked()  : save_formats.append('doc')
        if self.docOption_2.isChecked() : save_formats.append('ppt')
        if self.docOption_3.isChecked() : save_formats.append('pptx')
        converter = img2doc(path, save_to, save_formats=save_formats, font=font, language=language)       
        converter.convert()
        converter.merge()
        self.pushButton_3.setText('Done Converting!')
        
    def oca_path(self):
    	self.pushButton_3.setText('Convert')
    	self.pathOption.setText(QFileDialog.getExistingDirectory())
    
    def oca_save(self):
    	self.saveToOption.setText(QFileDialog.getExistingDirectory())
    	predicted_time = (2 * len(os.listdir(self.pathOption.text())))
    	time_type = 'seconds'
    	if predicted_time >= 60:
    		predicted_time = self.predicted_time // 60 
    		time_type = 'minute' if predicted_time == 1 else 'minutes'
    	self.pushButton_3.setText('takes '+' '+str(predicted_time)+' '+time_type)
if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    MainWindow = QtWidgets.QMainWindow()
    ui = Ui_MainWindow()
    ui.setupUi(MainWindow)
    MainWindow.show()
    sys.exit(app.exec_())
